#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate Comprehensive System Analysis Report in DOCX Format
Exam Paper & Assignment Generator - Full Analysis with ER Diagram & DFD
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# ========== TITLE PAGE ==========
title = doc.add_heading('EXAM PAPER & ASSIGNMENT GENERATOR', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_heading('System Analysis Report', 2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add metadata
meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Light Grid Accent 1'
meta_table.rows[0].cells[0].text = 'Project'
meta_table.rows[0].cells[1].text = 'AI Powered Exam Paper Generator'
meta_table.rows[1].cells[0].text = 'Location'
meta_table.rows[1].cells[1].text = 'Backend/AI_Powered_Exam_Paper_Generator/'
meta_table.rows[2].cells[0].text = 'Date'
meta_table.rows[2].cells[1].text = datetime.now().strftime("%d-%B-%Y")
meta_table.rows[3].cells[0].text = 'Report Type'
meta_table.rows[3].cells[1].text = 'Database Schema + Process Flow Analysis'

doc.add_paragraph()

# ========== TABLE OF CONTENTS ==========
doc.add_heading('TABLE OF CONTENTS', 1)
toc_items = [
    '1. Executive Summary',
    '2. System Overview',
    '3. Entity Relationship Diagram (ER Diagram)',
    '4. Level 1 Data Flow Diagram (DFD)',
    '5. Database Schema Details',
    '6. Process Descriptions',
    '7. Key Features & Algorithms',
    '8. API Endpoints',
    '9. Technical Stack',
    '10. Code Architecture',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ========== EXECUTIVE SUMMARY ==========
doc.add_heading('1. EXECUTIVE SUMMARY', 1)

doc.add_heading('Purpose', 2)
doc.add_paragraph(
    'The Exam Paper & Assignment Generator is an intelligent system that automatically '
    'generates exam papers using historical question data (PYQ - Previous Year Questions). '
    'It employs frequency-based ranking algorithms to select the most relevant questions '
    'and formats them into professional university-standard PDF documents.'
)

doc.add_heading('Key Capabilities', 2)
capabilities = [
    'Automated question selection from PYQ database',
    'Frequency-based ranking algorithm (by repetition, time span)',
    'Support for Internal (30 marks) and External (60 marks) exams',
    'Professional PDF generation in Ganpat University format',
    'Complete paper history tracking in database',
    'Question variation generation for insufficient data',
    'Multi-semester and multi-subject support',
]
for cap in capabilities:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_heading('System Scope', 2)
scope_table = doc.add_table(rows=5, cols=2)
scope_table.style = 'Light Grid Accent 1'
scope_table.rows[0].cells[0].text = 'Component'
scope_table.rows[0].cells[1].text = 'Details'
scope_table.rows[1].cells[0].text = 'Users'
scope_table.rows[1].cells[1].text = 'Faculty (paper generation), Students (paper download)'
scope_table.rows[2].cells[0].text = 'Semesters'
scope_table.rows[2].cells[1].text = '1-6 (configurable)'
scope_table.rows[3].cells[0].text = 'Exam Types'
scope_table.rows[3].cells[1].text = 'Internal (30 marks), External (60 marks)'
scope_table.rows[4].cells[0].text = 'Data Source'
scope_table.rows[4].cells[1].text = 'PYQ data (Previous Year Questions from CSV/JSON)'

doc.add_page_break()

# ========== SYSTEM OVERVIEW ==========
doc.add_heading('2. SYSTEM OVERVIEW', 1)

doc.add_heading('Architecture Overview', 2)
overview_img = doc.add_paragraph(
    'The system operates through 5 main processes:\n\n'
    '1. Faculty selects semester & subject\n'
    '2. System retrieves ranked questions from PYQ database\n'
    '3. Questions selected based on exam type and distribution\n'
    '4. PDF formatted in university standard layout\n'
    '5. Paper saved to database and delivered to user'
)

doc.add_heading('Key Input Sources', 2)
sources = [
    ('PYQ Historical Data', 'CSV/JSON files with previous year questions'),
    ('ExamPaperConfig', 'Configuration class with semester-subject mappings'),
    ('Faculty Request', 'Semester, subject code, exam type selection'),
    ('Question Bank', 'Templates for variations if insufficient data'),
]
for source, desc in sources:
    doc.add_paragraph(f'{source}: {desc}', style='List Bullet 2')

doc.add_heading('Key Output', 2)
doc.add_paragraph(
    'Professional PDF exam paper in Ganpat University format with:\n'
    '• University header and branding\n'
    '• Subject and course identification\n'
    '• Exam instructions\n'
    '• Properly formatted questions with marks\n'
    '• Time duration and constraints',
    style='List Bullet'
)

doc.add_page_break()

# ========== ER DIAGRAM ==========
doc.add_heading('3. ENTITY RELATIONSHIP DIAGRAM (ER DIAGRAM)', 1)

doc.add_paragraph(
    'The ER Diagram shows the complete database schema with 24 tables organized '
    'into 7 categories. The main focus is on exam generation, student results, '
    'and course scheduling.'
)

doc.add_heading('ER Diagram Structure', 2)

# ER Diagram table
er_table = doc.add_table(rows=8, cols=3)
er_table.style = 'Light Grid Accent 1'
hdr = er_table.rows[0].cells
hdr[0].text = 'Entity'
hdr[1].text = 'Primary Purpose'
hdr[2].text = 'Key Relationships'

data = [
    ('COURSE', 'Degree program definition', 'Has SUBJECT, TIMETABLE_SLOT'),
    ('SUBJECT', 'Course module per semester', 'Part of COURSE, Has EXAM, Generates GENERATED_PAPER'),
    ('GENERATED_PAPER', 'Core: Generated exam papers', 'Belongs to SUBJECT, Contains QUESTION selection'),
    ('EXAM', 'Exam configuration', 'Created by FACULTY, Belongs to SUBJECT, Contains QUESTION'),
    ('QUESTION', 'Individual exam questions', 'Belongs to EXAM, Varies in type (MCQ/Short/Long)'),
    ('STUDENT', 'Learner profile', 'Takes EXAM, Achieves SEMESTER_RESULT'),
    ('FACULTY', 'Instructor profile', 'Creates EXAM, Teaches in TIMETABLE_SLOT'),
]

for row_idx, (entity, purpose, relations) in enumerate(data, 1):
    row = er_table.rows[row_idx].cells
    row[0].text = entity
    row[1].text = purpose
    row[2].text = relations

doc.add_heading('Core Paper Generation Path', 2)
doc.add_paragraph(
    'FACULTY → Creates → EXAM ← Contains ← QUESTION\n'
    '         ↓\n'
    'EXAM belongs to SUBJECT (fk) ← belongs to ← COURSE\n'
    '         ↓\n'
    'GeneratedPaper stores selected QUESTION items in JSON format\n'
    'for specific SUBJECT + Semester',
    style='List Bullet'
)

doc.add_heading('ER Diagram - Mermaid Code', 2)
code_para = doc.add_paragraph(
    'Full ER diagram code available in: ER_DIAGRAM_MERMAID.md'
)
code_para.runs[0].font.italic = True

doc.add_page_break()

# ========== DFD ==========
doc.add_heading('4. LEVEL 1 DATA FLOW DIAGRAM (DFD)', 1)

doc.add_paragraph(
    'The Level 1 DFD illustrates the main process flow showing how data flows '
    'through 5 sequential processes from user input to PDF delivery.'
)

doc.add_heading('DFD Processes Overview', 2)

# DFD Process table
dfd_table = doc.add_table(rows=6, cols=4)
dfd_table.style = 'Light Grid Accent 1'
hdr = dfd_table.rows[0].cells
hdr[0].text = 'Process'
hdr[1].text = 'Input'
hdr[2].text = 'Processing'
hdr[3].text = 'Output'

dfd_data = [
    ('1.0 Select', 'Faculty action', 'Accept semester & subject', 'Parameters'),
    ('2.0 Retrieve', 'Sem + Code', 'Query PYQ, rank by frequency', 'Ranked questions'),
    ('3.0 Select', 'Ranked Qs', 'Filter by exam type & count', 'Selected questions'),
    ('4.0 Generate', 'Selected Qs', 'Create PDF with formatting', 'PDF file'),
    ('5.0 Store', 'PDF + Meta', 'Save to database', 'Stored record'),
]

for row_idx, (proc, inp, process, out) in enumerate(dfd_data, 1):
    row = dfd_table.rows[row_idx].cells
    row[0].text = proc
    row[1].text = inp
    row[2].text = process
    row[3].text = out

doc.add_heading('Data Stores in DFD', 2)
stores = [
    ('D1: PYQ Data', 'Historical question repository with frequency metrics'),
    ('D2: ExamPaperConfig', 'Subject code to name mappings per semester'),
    ('D3: GeneratedPaper', 'Database table storing all generated papers'),
    ('D4: Question Bank', 'Question variation templates for padding'),
]
for store, desc in stores:
    doc.add_paragraph(f'{store}: {desc}', style='List Bullet')

doc.add_heading('DFD - Mermaid Code', 2)
code_para = doc.add_paragraph(
    'Full DFD code available in: DFD_LEVEL1_MERMAID.md'
)
code_para.runs[0].font.italic = True

doc.add_page_break()

# ========== DATABASE SCHEMA ==========
doc.add_heading('5. DATABASE SCHEMA DETAILS', 1)

doc.add_heading('GeneratedPaper Table (Core)', 2)
gp_table = doc.add_table(rows=9, cols=3)
gp_table.style = 'Light Grid Accent 1'
hdr = gp_table.rows[0].cells
hdr[0].text = 'Column'
hdr[1].text = 'Type'
hdr[2].text = 'Purpose'

gp_data = [
    ('id', 'INT PK', 'Primary key (auto-increment)'),
    ('semester', 'INT', 'Semester number (1-6)'),
    ('subject_code', 'VARCHAR', 'Subject code (e.g., CS101)'),
    ('subject_name', 'VARCHAR', 'Subject name (e.g., Data Structures)'),
    ('exam_type', 'VARCHAR', 'Internal or External'),
    ('total_marks', 'INT', '30 (internal) or 60 (external)'),
    ('paper_data', 'JSON', 'Complete question structure'),
    ('created_at', 'DATETIME', 'Generation timestamp'),
]

for row_idx, (col, typ, purpose) in enumerate(gp_data, 1):
    row = gp_table.rows[row_idx].cells
    row[0].text = col
    row[1].text = typ
    row[2].text = purpose

doc.add_heading('Question Table (academics app)', 2)
doc.add_paragraph(
    'Structure: question_id (PK), exam_id (FK), question_text, question_type, marks, options (JSON)\n'
    'Types: MCQ, Short Answer, Long Answer'
)

doc.add_heading('Exam Table (academics app)', 2)
doc.add_paragraph(
    'Structure: exam_id (PK), subject_id (FK), exam_type, date, duration, total_marks, created_by (Faculty FK)\n'
    'Exam Types: Mid Term, End Term, Quiz, Assignment, Practical'
)

doc.add_page_break()

# ========== PROCESS DESCRIPTIONS ==========
doc.add_heading('6. PROCESS DESCRIPTIONS', 1)

doc.add_heading('Process 1.0: Select Semester & Subject', 2)
doc.add_paragraph('Faculty selects semester (1-6) and subject code from dropdowns.')
doc.add_paragraph('API: GET /exam_paper/get_semesters/ and GET /exam_paper/get_subjects/{semester}/')

doc.add_heading('Process 2.0: Retrieve & Rank Questions', 2)
doc.add_paragraph(
    'System queries PYQ data and sorts by:\n'
    '• year_count (appeared in more years = higher priority)\n'
    '• count (total frequency}\n'
    '• freq_score (computed importance)',
    style='List Bullet'
)

doc.add_heading('Process 3.0: Question Selection', 2)
doc.add_paragraph('Algorithm based on exam type:')
doc.add_paragraph(
    'External (60 marks, 3 hours):\n'
    '• Section 1: 9 MCQs\n'
    '• Q2A: 2 medium questions\n'
    '• Q2B: 2 medium questions\n'
    '• Q3: 2 long answer\n'
    '• Q4: 3 complex problem-solving\n'
    '• Q5: 2 difficult application questions',
    style='List Bullet'
)
doc.add_paragraph(
    'Internal (30 marks, 1.5 hours):\n'
    '• Section 1: 9 MCQs (light assessment)',
    style='List Bullet'
)

doc.add_heading('Process 4.0: Generate PDF', 2)
doc.add_paragraph(
    'Creates professional PDF with:\n'
    '• Ganpat University header\n'
    '• Course & subject details\n'
    '• Standard instructions\n'
    '• Formatted questions with marks\n'
    '• Time & total marks indication',
    style='List Bullet'
)

doc.add_heading('Process 5.0: Store & Deliver', 2)
doc.add_paragraph(
    'Saves paper record to GeneratedPaper table and delivers PDF to user as download.'
)

doc.add_page_break()

# ========== KEY FEATURES ==========
doc.add_heading('7. KEY FEATURES & ALGORITHMS', 1)

doc.add_heading('1. Intelligent Question Ranking', 2)
doc.add_paragraph(
    'Questions ranked by historical frequency metrics:\n'
    '- How often a question appeared\n'
    '- Over how many years it appeared\n'
    '- Combined importance score\n\n'
    'Most frequently repeated questions get priority for exam papers.'
)

doc.add_heading('2. Question Variation Generation', 2)
doc.add_paragraph(
    'If insufficient unique questions exist, system generates variations:\n\n'
    'Variation templates:\n'
    '- "Explain in detail: [base question]"\n'
    '- "With suitable example, describe: [base]"\n'
    '- "Write short note on: [base]"\n'
    '- "Critically analyze: [base]"\n'
    '- And 5 more template variations\n\n'
    'Ensures papering requirements always met even with limited data.'
)

doc.add_heading('3. Exam-Type Aware Selection', 2)
doc.add_paragraph(
    'Different algorithms for Internal vs External exams:\n'
    '- Internal: Lighter assessment, only MCQs\n'
    '- External: Comprehensive coverage with multiple question types\n'
    '- Both respect time and marks constraints'
)

doc.add_heading('4. University-Standard PDF Format', 2)
doc.add_paragraph(
    'Generated PDFs follow exact Ganpat University format:\n'
    '- Enrollment number field\n'
    '- University branding\n'
    '- Course, semester, exam type header\n'
    '- Standard instructions\n'
    '- Professional typography'
)

doc.add_heading('5. Complete Paper History', 2)
doc.add_paragraph(
    'Every generated paper stored in database:\n'
    '- Timestamp of generation\n'
    '- Examination metadata\n'
    '- Complete question structure in JSON\n'
    '- Enables re-generation and analytics'
)

doc.add_page_break()

# ========== API ENDPOINTS ==========
doc.add_heading('8. API ENDPOINTS', 1)

apis = [
    ('GET /exam_paper/get_semesters/', 'Returns list of available semesters [1, 2, 3, 4, 5, 6]'),
    ('GET /exam_paper/get_subjects/{semester}/', 'Returns subjects for semester: [{code, name}, ...]'),
    ('POST /exam_paper/generate_paper/', 'Generates paper: {semester, subject_code, exam_type}'),
    ('POST /exam_paper/download_pdf/', 'Creates PDF, saves DB record, returns file'),
]

for endpoint, desc in apis:
    doc.add_paragraph(f'{endpoint}\n→ {desc}', style='List Bullet')

doc.add_page_break()

# ========== TECHNICAL STACK ==========
doc.add_heading('9. TECHNICAL STACK', 1)

stack_table = doc.add_table(rows=7, cols=2)
stack_table.style = 'Light Grid Accent 1'
stack_table.rows[0].cells[0].text = 'Layer'
stack_table.rows[0].cells[1].text = 'Technology'
stack_table.rows[1].cells[0].text = 'Backend'
stack_table.rows[1].cells[1].text = 'Django 4.x + Django REST Framework'
stack_table.rows[2].cells[0].text = 'PDF Generation'
stack_table.rows[2].cells[1].text = 'FPDF library (pure Python)'
stack_table.rows[3].cells[0].text = 'Data Processing'
stack_table.rows[3].cells[1].text = 'Frequency ranking algorithm, JSON processing'
stack_table.rows[4].cells[0].text = 'Database'
stack_table.rows[4].cells[1].text = 'PostgreSQL with JSONField'
stack_table.rows[5].cells[0].text = 'Configuration'
stack_table.rows[5].cells[1].text = 'ExamPaperConfig class (app config)'
stack_table.rows[6].cells[0].text = 'Data Source'
stack_table.rows[6].cells[1].text = 'PYQ data (CSV/JSON importable)'

doc.add_page_break()

# ========== CODE ARCHITECTURE ==========
doc.add_heading('10. CODE ARCHITECTURE', 1)

doc.add_heading('File Structure', 2)
doc.add_paragraph(
    'Backend/AI_Powered_Exam_Paper_Generator/\n'
    '├── views.py (HTTP endpoints)\n'
    '├── utils.py (Question ranking, PDF generation)\n'
    '├── models.py (GeneratedPaper table)\n'
    '├── urls.py (Route definitions)\n'
    '├── apps.py (ExamPaperConfig with question data)\n'
    '├── migrations/ (Database schema)\n'
    '└── static/\n'
    '    └── templates/'
)

doc.add_heading('Key Functions', 2)

functions = [
    ('get_ranked_questions(sem, code)', 'Retrieves & ranks PYQ questions'),
    ('pick_questions(ranked, n, used)', 'Selects n unique questions, generates variations'),
    ('create_pdf(paper)', 'Generates PDF from question structure'),
    ('generate_paper(request)', 'Main endpoint: orchestrates 1.0-3.0 processes'),
    ('download_pdf(request)', 'Creates PDF and stores in DB (processes 4.0-5.0)'),
]

for func, desc in functions:
    doc.add_paragraph(f'{func}\n→ {desc}', style='List Bullet 2')

doc.add_heading('Data Flow in Code', 2)
doc.add_paragraph(
    '1. Faculty POST request to /exam_paper/generate_paper/\n'
    '2. generate_paper() calls get_ranked_questions()\n'
    '3. get_ranked_questions() loads ExamPaperConfig.clusters\n'
    '4. pick_questions() selects based on exam_type\n'
    '5. Faculty downloads, triggers download_pdf()\n'
    '6. create_pdf() generates FPDF instance\n'
    '7. GeneratedPaper record created in database\n'
    '8. FileResponse returns PDF to browser'
)

# ========== FOOTER ==========
doc.add_page_break()
doc.add_heading('APPENDIX', 1)

doc.add_heading('References', 2)
references = [
    'Full ER Diagram Mermaid code: ER_DIAGRAM_MERMAID.md',
    'Full DFD Level 1 Mermaid code: DFD_LEVEL1_MERMAID.md',
    'Source code: Backend/AI_Powered_Exam_Paper_Generator/',
    'Database models: Backend/academics/models.py',
]
for ref in references:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('Contact & Support', 2)
doc.add_paragraph('For questions about this analysis, refer to project documentation or contact the development team.')

# Footer
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.add_run('\n' + '='*60 + '\n')
footer_para.add_run('Generated by System Analysis Tool | AMPICS Academic Module\n')
footer_para.add_run(f'Report Date: {datetime.now().strftime("%d-%B-%Y at %I:%M %p")}\n')
footer_para.runs[0].font.size = Pt(9)
footer_para.runs[0].font.italic = True

# Save document
output_path = r'c:\Academic-module\EXAM_GENERATOR_ANALYSIS_REPORT.docx'
doc.save(output_path)
print(f'✅ Analysis Report created: {output_path}')

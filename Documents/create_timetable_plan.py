#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate Timetable Plan Document in DOCX Format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Add title
title = doc.add_heading('TIMETABLE GENERATION PLAN 📅', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add date
date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%d-%B-%Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para_format = date_para.runs[0].font
date_para_format.italic = True
date_para_format.size = Pt(11)

doc.add_paragraph()  # Spacing

# ==================== EXECUTIVE SUMMARY ====================
doc.add_heading('EXECUTIVE SUMMARY', 1)
summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = 'Light Grid Accent 1'
hdr_cells = summary_table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Status'

rows_data = [
    ('Faculty Assignments', '✅ Ready - Assigned to subjects/courses'),
    ('Database Schema', '✅ Ready - TimetableSlot models exist'),
    ('Generation Engine', '⚠️ Partial - Saturday limited to 3 (need 4)'),
    ('Room Management', '✅ Ready - Capacity & types configured'),
    ('Monday-Friday', '✅ Full time slots available'),
]

for component, status in rows_data:
    row_cells = summary_table.add_row().cells
    row_cells[0].text = component
    row_cells[1].text = status

doc.add_paragraph()

# ==================== REQUIREMENTS ====================
doc.add_heading('📋 REQUIREMENTS', 1)
req_list = [
    'Monday-Friday: FULL time table with complete slot coverage',
    'Saturday: EXACTLY 4 lectures (Current: 3 - needs update)',
    'Faculty already assigned to subjects & courses',
    'No conflicts: Faculty, Room, or Time overlaps',
]

for req in req_list:
    doc.add_paragraph(req, style='List Number')

doc.add_paragraph()

# ==================== ALGORITHM STRATEGY ====================
doc.add_heading('🎯 BEST ALGORITHM APPROACH', 1)

doc.add_heading('5-Phase Greedy Algorithm with Backtracking', 2)

phases = {
    'PHASE 1: Preparation (Data Validation)': [
        'Verify all Faculty → Subject mappings',
        'Validate Room availability & capacity',
        'Check TimeSlot configuration for both shifts',
        'Ensure no circular dependencies',
    ],
    'PHASE 2: Constraint Definition': [
        'HARD CONSTRAINTS:',
        '  • Faculty cannot teach multiple classes simultaneously',
        '  • Room cannot have overlapping bookings',
        '  • Teacher max_lectures_per_day must be respected',
        '  • Campus branch consistency (Kherva ≠ Ahmedabad)',
        '',
        'SOFT CONSTRAINTS:',
        '  • Faculty working_shift preference',
        '  • Subject teaching preference (Theory/Practical)',
        '  • Student load balancing',
    ],
    'PHASE 3: Slot Generation': [
        'WEEKDAY (Monday-Friday):',
        '  • MORNING: 8:00-12:55 (5 hours × multiple slots)',
        '  • NOON: 12:00-18:10 (6 hours × multiple slots)',
        '  • Distribute subjects evenly across days',
        '',
        'SATURDAY:',
        '  • EXACTLY 4 lectures (not 3)',
        '  • Light schedule for specialized subjects',
    ],
    'PHASE 4: Assignment (Greedy with Backtracking)': [
        'For each Subject in each Course:',
        '  1. Get assigned Faculty pool',
        '  2. For each required Theory/Practical session:',
        '     a. Pick optimal Faculty (lowest current load)',
        '     b. Find next available time slot (Mon-Fri or Sat)',
        '     c. Book suitable Room (by capacity & branch)',
        '     d. Record TimetableSlot entry',
        '     e. If conflict detected: Try next faculty or time',
    ],
    'PHASE 5: Optimization': [
        'Balance faculty workload evenly',
        'Minimize gaps between consecutive classes',
        'Respect lunch/prayer/tea breaks',
        'Post-process conflict resolution',
    ],
}

for phase, details in phases.items():
    doc.add_heading(phase, 3)
    for detail in details:
        if detail.startswith('  •'):
            p = doc.add_paragraph(detail, style='List Bullet 2')
        elif detail.startswith(' '):
            p = doc.add_paragraph(detail)
        else:
            p = doc.add_paragraph(detail, style='List Bullet')

doc.add_paragraph()

# ==================== DATABASE CONFIGURATION ====================
doc.add_heading('📊 DATABASE CONFIGURATION', 1)

doc.add_heading('Current TimeSlot Setup', 2)
doc.add_paragraph('Expected configuration in database:', style='List Bullet')

config_list = [
    'Morning Shift: 5 slots (08:00, 08:55, 10:15, 11:10, 12:00)',
    'Noon Shift: 5 slots (12:00, 13:25, 15:15, 16:30, 17:20)',
    'Saturday: Need 4 slots (currently limited to 3)',
]

for config in config_list:
    doc.add_paragraph(config, style='List Bullet 2')

doc.add_heading('Validation Query', 2)
code_para = doc.add_paragraph()
code_para.add_run('SELECT * FROM time_slots ').font.name = 'Courier New'
code_para.add_run('WHERE campus_branch').font.name = 'Courier New'
code_para.add_run("='Kherva' AND day_type IS NULL;").font.name = 'Courier New'

doc.add_paragraph()

# ==================== CRITICAL FIXES ====================
doc.add_heading('🔧 CRITICAL CONFIGURATION CHANGES', 1)

doc.add_heading('Change #1: Update Saturday Slots from 3 to 4', 2)
doc.add_paragraph('📁 File: Backend/academics/management/commands/generate_timetable.py')
doc.add_paragraph('📍 Line: 37')

# Before/After code
before_para = doc.add_paragraph()
before_para.add_run('❌ CURRENT (Wrong):').font.bold = True
before_code = doc.add_paragraph('SATURDAY_SLOT_LIMIT = 3', style='List Bullet')
before_code.paragraph_format.left_indent = Inches(0.5)

after_para = doc.add_paragraph()
after_para.add_run('✅ REQUIRED (Correct):').font.bold = True
after_code = doc.add_paragraph('SATURDAY_SLOT_LIMIT = 4', style='List Bullet')
after_code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Change #2: Verify Monday-Friday Slots', 2)
weekday_para = doc.add_paragraph('✅ Status: Already configured correctly')
weekday_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph('Morning: 5 slots providing full coverage', style='List Bullet 2')
doc.add_paragraph('Noon: 5 slots providing full coverage', style='List Bullet 2')

doc.add_paragraph()

# ==================== IMPLEMENTATION STEPS ====================
doc.add_heading('🚀 STEP-BY-STEP IMPLEMENTATION', 1)

steps = {
    'STEP 1: Update Saturday Configuration': [
        'Navigate to: Backend/academics/management/commands/generate_timetable.py',
        'Find line 37: SATURDAY_SLOT_LIMIT = 3',
        'Change to: SATURDAY_SLOT_LIMIT = 4',
        'Save the file',
    ],
    'STEP 2: Validate Data in Database': [
        'Open Django shell: python manage.py shell',
        'Run validation checks:',
        '  from academics.models import Course, Subject, TimeSlot',
        '  print("Courses:", Course.objects.count())',
        '  print("Subjects:", Subject.objects.count())',
        '  print("TimeSlots:", TimeSlot.objects.count())',
    ],
    'STEP 3: Pre-Generation Checklist': [
        '☑ All Faculty assigned to subjects',
        '☑ TimeSlot table populated (morning/noon/saturday)',
        '☑ Rooms created with adequate capacity',
        '☑ Course shift properly set (MORNING/NOON)',
        '☑ Saturday config updated (4 lectures)',
        '☑ Database backup taken (recommended)',
    ],
    'STEP 4: Generate Timetable': [
        'Option A - Django Command:',
        '  cd Backend',
        '  python manage.py generate_timetable',
        '',
        'Option B - Via API (requires admin authentication):',
        '  POST /academics/admin/timetable/generate-complete/',
    ],
    'STEP 5: Verify Results': [
        'Check generation status:',
        '  GET /academics/admin/timetable/stats/',
        '',
        'View all timetable slots:',
        '  GET /academics/timetable/',
        '',
        'Verify Saturday has 4 lectures per course',
        'Verify Monday-Friday has full coverage',
    ],
}

for step, actions in steps.items():
    doc.add_heading(step, 2)
    for idx, action in enumerate(actions, 1):
        if action.startswith('  '):
            p = doc.add_paragraph(action)
            p.paragraph_format.left_indent = Inches(0.5)
        elif '.' in action and idx == 1:
            doc.add_paragraph(action, style='List Bullet')
        else:
            doc.add_paragraph(action, style='List Bullet')

doc.add_paragraph()

# ==================== POST-GENERATION ====================
doc.add_heading('⚡ POST-GENERATION ADJUSTMENTS', 1)

doc.add_paragraph('After successful generation, you can make manual adjustments:')

adjustments = [
    'Lock specific slots to prevent regeneration: is_locked=True',
    'Edit faculty/room assignments for exceptional cases',
    'Add manual slots for special lectures or replacement sessions',
    'Generate detailed conflict reports',
    'Export timetable to Excel/PDF format for printing',
]

for adj in adjustments:
    doc.add_paragraph(adj, style='List Bullet')

doc.add_paragraph()

# ==================== EXPECTED TIMETABLE STRUCTURE ====================
doc.add_heading('📅 EXPECTED TIMETABLE STRUCTURE', 1)

# Monday-Friday table
weekday_table = doc.add_table(rows=6, cols=3)
weekday_table.style = 'Light Grid Accent 1'

# Header
hdr_cells = weekday_table.rows[0].cells
hdr_cells[0].text = 'Time Slot'
hdr_cells[1].text = 'Morning Shift'
hdr_cells[2].text = 'Noon Shift'

# Morning slots
morning_slots = [
    ('08:00 - 08:55', 'Slot 1', 'Slot 1'),
    ('08:55 - 09:40', 'Slot 2', 'Slot 2'),
    ('10:15 - 11:10', 'Slot 3', 'Slot 3'),
    ('11:10 - 12:00', 'Slot 4', 'Slot 4'),
    ('12:00 - 12:55', 'Slot 5', 'Slot 5'),
]

for idx, (time, morning, noon) in enumerate(morning_slots, 1):
    row_cells = weekday_table.rows[idx].cells
    row_cells[0].text = time
    row_cells[1].text = morning
    row_cells[2].text = noon

doc.add_paragraph('Monday to Friday: Each slot has Subject+Faculty+Room', style='List Bullet')

doc.add_paragraph()

# Saturday schedule
sat_table = doc.add_table(rows=5, cols=2)
sat_table.style = 'Light Grid Accent 1'

sat_hdr = sat_table.rows[0].cells
sat_hdr[0].text = 'Saturday Lecture'
sat_hdr[1].text = 'Duration'

sat_data = [
    ('Lecture 1', '60 minutes'),
    ('Lecture 2', '60 minutes'),
    ('Lecture 3', '60 minutes'),
    ('Lecture 4', '60 minutes'),
]

for idx, (lecture, duration) in enumerate(sat_data, 1):
    row_cells = sat_table.rows[idx].cells
    row_cells[0].text = lecture
    row_cells[1].text = duration

doc.add_paragraph('Saturday: 4 lectures for selected subjects', style='List Bullet')

doc.add_paragraph()

# ==================== TROUBLESHOOTING ====================
doc.add_heading('❓ TROUBLESHOOTING', 1)

troubleshoots = {
    'Problem: Saturday still showing 3 lectures': [
        '✓ Check that SATURDAY_SLOT_LIMIT = 4 was updated',
        '✓ Clear existing timetable slots before regenerating',
        '✓ Restart Django server after code changes',
    ],
    'Problem: Faculty conflicts detected': [
        '✓ Verify faculty working_shift matches course shift',
        '✓ Check max_lectures_per_day is set correctly',
        '✓ Ensure adequate faculty pool for subject coverage',
    ],
    'Problem: Room booking conflicts': [
        '✓ Verify all rooms have is_available=True',
        '✓ Check room capacity >= student enrollment',
        '✓ Ensure campus_branch consistency',
    ],
    'Problem: Generation takes too long': [
        '✓ Reduce number of courses in test run',
        '✓ Check database indexes on timetable tables',
        '✓ Verify no circular dependencies in data',
    ],
}

for problem, solutions in troubleshoots.items():
    doc.add_heading(problem, 2)
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

doc.add_paragraph()

# ==================== KEY FILES REFERENCE ====================
doc.add_heading('📁 KEY FILES REFERENCE', 1)

files_ref = {
    'Generation Engine': 'Backend/academics/management/commands/generate_timetable.py',
    'Models': 'Backend/academics/models.py (TimetableSlot, Course, Subject)',
    'Views/APIs': 'Backend/academics/views.py (admin_generate_timetable)',
    'URLs': 'Backend/academics/urls.py',
    'Serializers': 'Backend/academics/serializers.py (TimetableSlotSerializer)',
}

for component, path in files_ref.items():
    doc.add_paragraph(f'{component}: {path}', style='List Bullet')

doc.add_paragraph()

# ==================== QUICK COMMAND REFERENCE ====================
doc.add_heading('⚙️ QUICK COMMAND REFERENCE', 1)

commands = [
    ('Generate Timetable', 'python manage.py generate_timetable'),
    ('Check Stats', 'GET /academics/admin/timetable/stats/'),
    ('View All Slots', 'GET /academics/timetable/'),
    ('Clear Timetable', 'GET /academics/admin/timetable/clear/'),
    ('Export to PDF', 'GET /academics/admin/timetable/pdf/'),
]

cmd_table = doc.add_table(rows=1, cols=2)
cmd_table.style = 'Light Grid Accent 1'

hdr = cmd_table.rows[0].cells
hdr[0].text = 'Action'
hdr[1].text = 'Command'

for action, cmd in commands:
    row = cmd_table.add_row().cells
    row[0].text = action
    row[1].text = cmd

doc.add_paragraph()

# ==================== FOOTER ====================
doc.add_paragraph()
footer_para = doc.add_paragraph('Generated by Timetable Generation System | AMPICS Academic Module')
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.runs[0].font.size = Pt(10)
footer_para.runs[0].font.italic = True

# Save document
output_path = r'c:\Academic-module\TIMETABLE_GENERATION_PLAN.docx'
doc.save(output_path)
print(f'✅ Document created successfully: {output_path}')
